"""
Subject: Scrape info of Canadian parliament members and generate reports in .xlsx and .docx (without using pandas)
Developer: Magg Lui
Last Modified: 2023-07-15
"""
import os
import requests
from datetime import date
from bs4 import BeautifulSoup
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.drawing.image import Image
from PIL import Image as PIL_Image
from docx import Document


# Create a function to find the member count by a specific key
def count_member(key: str, dict_list: list, sort_keys=True):
    """
    :param key: String, a key in a dictionary
    :param dict_list: List, a list of dictionaries, ideally the dictionaries share the same set of keys
    :return: Dictionary, where the keys are the unique values associated with the key specified in the param,
    and the values are the occurrence counts; the dictionary is sorted by the key by default, otherwise it is sorted in
    descending order by the values
    """
    unique_values = sorted(set(list(dictionary[key] for dictionary in dict_list)))
    results = {}
    for unique_value in unique_values:
        count = 0
        for dictionary in dict_list:
            if dictionary[key] == unique_value:
                count += 1
        results[unique_value] = count
    if sort_keys:
        return results
    else:
        results = dict(sorted(results.items(), key=lambda x: x[1], reverse=True))
    return results


# Create another function to find the member count by a pair of keys
def cross_count_member(key1: str, key2: str, dict_list: list, sort_keys=True):
    """
    :param key1: String, a key in a dictionary
    :param key2: String, another key in the dictionary
    :param dict_list: List, a list of dictionaries, ideally the dictionaries share the same set of keys
    :return: Dictionary, where the keys are the unique values associated with key1 (and sorted accordingly),
             and the values are dictionaries in which the keys are the unique values associated with key2 (and sorted
             accordingly by default, otherwise sorted by values in descending order), and values are the occurrence counts
    """
    unique_values_1 = sorted(set(list(dictionary[key1] for dictionary in dict_list)))
    unique_values_2 = sorted(set(list(dictionary[key2] for dictionary in dict_list)))
    results = {}
    for unique_value_1 in unique_values_1:
        cross_count = {}
        for unique_value_2 in unique_values_2:
            count = 0
            for dictionary in dict_list:
                if dictionary[key2] == unique_value_2 and dictionary[key1] == unique_value_1:
                    count += 1
            if count > 0:
                cross_count[unique_value_2] = count
            if not sort_keys:
                cross_count = dict(sorted(cross_count.items(), key=lambda x:x[1], reverse=True))
            results[unique_value_1] = cross_count
    return results

# Create a third function to get the unique count of one key by another key
def cross_distinct_key(key1: str, key2: str, dict_list: list):
    """
    :param key1: String, a key in a dictionary
    :param key2: String, another key in the dictionary
    :param dict_list: List, a list of dictionaries, ideally the dictionaries share the same set of keys
    :return: Dictionary, where the keys are the unique values associated with key1 (and sorted accordingly),
             and the values are lists of key2 occurrence (also sorted)
    """
    unique_values_1 = sorted(set(list(dictionary[key1] for dictionary in dict_list)))
    results = {}
    for unique_value_1 in unique_values_1:
        distinct_key2 = set()
        for dictionary in dict_list:
            if dictionary[key1] == unique_value_1:
                distinct_key2.add(dictionary[key2])
        results[unique_value_1] = sorted(distinct_key2)
    return results


# Note the date for extracting info from the webpage
today = date.today()

# Create a new folder (if it does not already exist) in the current directory to store the output files
cwd = os.getcwd()
newDir = "parliament"
if not os.path.exists(newDir):
    os.mkdir(newDir)
    print(f"Created \'{newDir}\' directory in \'{cwd}\'.")
else:
    print(f"\'{newDir}\' directory already exists in \'{cwd}\'.")

# Change the working directory to the new folder
os.chdir(newDir)


# The base URL is the website's index url
baseURL = "https://www.ourcommons.ca"

# Scrape the web contents using BeautifulSoup
url = "https://www.ourcommons.ca/Members/en/search"
req = requests.get(url)
if req.status_code == 200:
    soup = BeautifulSoup(req.text, "lxml")

    # Each parliament member is contained in a pair of anchor tag <a></a> with class name "ce-mip-mp-tile"
    memberContainers = soup.find_all("a", {"class": "ce-mip-mp-tile"})

    # Create an empty list to store each member's info later
    members = []

    for memberContainer in memberContainers:
        # convert the relative path to absolute by adding the base URL
        memberURL = baseURL + memberContainer["href"]
        # ID is enclosed in brackets in the memberURL
        ID = memberURL[(memberURL.find("(")+1):-1]
        # full name is found in the div with class name "ce-mip-mp-name"
        fullname = memberContainer.find("div", {"class": "ce-mip-mp-name"}).text
        # party is found in the div with class name "ce-mip-mp-party"
        party = memberContainer.find("div", {"class": "ce-mip-mp-party"}).text
        # constituency is found in the div with class name "ce-mip-mp-constituency"
        constituency = memberContainer.find("div", {"class": "ce-mip-mp-constituency"}).text
        # province is found in the div with class name "ce-mip-mp-province"
        province = memberContainer.find("div", {"class": "ce-mip-mp-province"}).text
        # convert the relative path to absolute by adding the base URL
        imgURL = baseURL + memberContainer.img["src"]

        # store the member info in a dict and append to the members list
        member = {
            "ID": ID,
            "fullname": fullname,
            "party": party,
            "constituency": constituency,
            "province": province,
            "imgURL": imgURL,
            "memberURL": memberURL
        }
        members.append(member)

    print("Web contents scraped.")


    # Now output the results
    # Create a new workbook to save the members' info to an Excel file
    print("\nExporting web contents to .xlsx file...")
    wb = Workbook()
    ws = wb.active
    ws.title = "Parliament Members " + str(today)
    print(f"Active worksheet '{ws.title}'.")
    
    # Add column headers
    headers = ["ID", "Full Name", "Political Party", "Constituency", "Province", "Image", "Image URL", "Member Page URL"]
    ws.append(headers)
    print("Column headers added.")
    
    # Add rows
    print("Inserting entries...")
    for member in members:
        # The 6th column (column F) is left blank as the image files will be inserted later
        row = [member["ID"], member["fullname"], member["party"], member["constituency"], member["province"],
               "", member["imgURL"], member["memberURL"]]
        ws.append(row)
    
        # Adjust the column width based on the cell contents
        for column in ws.columns:
            column_letter = column[0].column_letter
            max_width = 0
            for cell in column:
                # for each column, the max column width will be based on the length of the longest cell value
                if len(cell.value) > max_width:
                    max_width = len(cell.value)
            # set the column width to be +1 to add buffer for higher readability
            ws.column_dimensions[column_letter].width = max_width + 1
    
        # download the member image (.jpg) locally and save the filename as the member's fullname
        imgFile = member["fullname"].replace(" ", "-") + ".jpg"
        with open(imgFile, "wb") as f:
            f.write(requests.get(member["imgURL"]).content)
    
        # insert the image to column F
        # row no. would be + 2 since the members list index starts at 0 and there's a header row
        img = openpyxl.drawing.image.Image(imgFile)
        rowNum = members.index(member) + 2
        img.anchor = "F" + str(rowNum)
        ws.add_image(img)
    
        # use the PIL Image module to get the size (in pixels) of the image
        imgWidth, imgHeight = PIL_Image.open(imgFile).size
        # convert the pixels to point units (see https://stackoverflow.com/a/70203371)
        # +1 is used to add buffer for higher readability
        ws.column_dimensions["F"].width = imgWidth / 7 + 1
        ws.row_dimensions[rowNum].height = imgHeight * 3/4 + 1
    
    print("All members added to worksheet.")
    
    # Bold the headers
    for cell in ws["1:1"]:
        cell.font = Font(bold=True)
    
    # Save the workbook
    excelFilename = f"members-as-of-{today}.xlsx"
    wb.save(filename=excelFilename)
    print(f"{excelFilename} created.")


    # Create a new docx for the analysis report
    print("\nGenerating docx report...")
    doc = Document()
    
    # Add heading and centre align the text
    doc.add_heading("Parliament Composition", 0).alignment = 1
    
    # Overview of results
    doc.add_paragraph(f"Date: {today}")
    totalCount = 0
    for f in os.listdir(os.getcwd()):
        if os.path.splitext(f)[1].lower() == ".jpg":
            totalCount += 1
    doc.add_paragraph(f"Total count of members: {totalCount}")
    
    # Add the section for member count by political party
    doc.add_heading("Member Count by Political Party", 1)
    partyTbl = doc.add_table(rows=1, cols=2)
    partyTblHeader = partyTbl.rows[0].cells
    partyTblHeader[0].paragraphs[0].add_run("Political Party").bold = True
    partyTblHeader[1].paragraphs[0].add_run("Member Count").bold = True
    partyCount = count_member("party", members, False)
    partyTblRows = tuple(pair for pair in partyCount.items())
    for party, count in partyTblRows:
        row_cell = partyTbl.add_row().cells
        row_cell[0].text = party
        row_cell[1].text = str(count)
    doc.add_paragraph("")

    # Add the section for member count by provinces
    doc.add_heading("Member Count by Province", 1)
    provinceTbl = doc.add_table(rows=1, cols=2)
    provinceTblHeader = provinceTbl.rows[0].cells
    provinceTblHeader[0].paragraphs[0].add_run("Province").bold = True
    provinceTblHeader[1].paragraphs[0].add_run("Member Count").bold = True
    provinceCount = count_member("province", members)
    provinceTblRows = tuple(pair for pair in provinceCount.items())
    for province, count in provinceTblRows:
        row_cell = provinceTbl.add_row().cells
        row_cell[0].text = province
        row_cell[1].text = str(count)
    doc.add_paragraph("")

    # Add the section for party member count by province
    doc.add_heading("Party Member Count by Province", 1)
    provincePartyCount = cross_count_member("province", "party", members, False)
    for province, partyDict in provincePartyCount.items():
        doc.add_heading(province, 2)
        provincePartyTbl = doc.add_table(rows=1, cols=2)
        provincePartyTblHeader = provincePartyTbl.rows[0].cells
        provincePartyTblHeader[0].paragraphs[0].add_run("Political Party").bold = True
        provincePartyTblHeader[1].paragraphs[0].add_run("Member Count").bold = True
        provincePartyTblRows = []
        for party, count in partyDict.items():
            provincePartyTblRows.append((party, count))
        for partyTuple in provincePartyTblRows:
            row_cell = provincePartyTbl.add_row().cells
            row_cell[0].text = partyTuple[0]
            row_cell[1].text = str(partyTuple[1])
        doc.add_paragraph("")


    # Add the section for provincial scope by party
    doc.add_heading("Provincial Scope by Party", 1)
    partyProvinceCount = cross_distinct_key("party", "province", members)
    for party in partyProvinceCount.keys():
        doc.add_heading(party, 2)
        partyProvinceTbl = doc.add_table(rows=0, cols=1)
        partyProvinceTblRows = list(province for province in partyProvinceCount[party])
        for province in partyProvinceTblRows:
            row_cell = partyProvinceTbl.add_row().cells
            row_cell[0].text = province
        doc.add_paragraph("")

    
    # Save the document
    docFilename = f"parliament-report-{today}.docx"
    doc.save(docFilename)
    print(f"{docFilename} created.")


# print an error msg if BeautifulSoup fails to scrape the contents from the parliament page
else:
    print("Something went wrong. Cannot scrape the web contents.")
